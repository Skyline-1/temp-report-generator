from io import BytesIO
import PIL.Image as PilImage
from scipy.interpolate import make_interp_spline
import anvil.server
import requests
import pandas as pd
import os
import statistics
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
import time
from datetime import datetime
from io import StringIO
import statistics
#import win32com.client
import math
import numpy as np
from anvil import BlobMedia 
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from anvil import *
import random

generated_ids = set()
def generate_unique_id():
    while True:
        new_id = str(random.randint(10000000, 99999999))
        if new_id not in generated_ids:
            generated_ids.add(new_id)
            return new_id
  
def get_temp_range(folder_path, set_point):
    if set_point == 20:
        return 15, 25
    elif set_point == 5:
        return 2, 8
    else:
        return -1, -1   
      
def update_heading_style(heading):
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
def update_summary_table_style(table, row_number):
    for cell in table.rows[row_number].cells:
    # Set background color to black
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '000000')  # Black color
        cell._element.get_or_add_tcPr().append(shading_elm)
    # Set text color to white
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
def update_style(table):
    for cell in table.rows[0].cells:
    # Set background color to black
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '000000')  # Black color
        cell._element.get_or_add_tcPr().append(shading_elm)
    # Set text color to white
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
def extract_temperatures_from_csv(file_path, start_datetime, end_datetime):
    temperatures = []
    print("file_path=", file_path)
    # Read the CSV file into a DataFrame
    bytes_data = file_path.get_bytes()
    df = pd.read_csv(BytesIO(bytes_data))
    print(df)
    df['DateTime'] = pd.to_datetime(df['Date'] + ' ' + df['Time'])
    filtered_df = df[(df['DateTime'] >= start_datetime) & (df['DateTime'] <= end_datetime)]
    # Extract the temperature values from the 'Temperature' column if it exists
    if 'Temperature' in filtered_df.columns:
        temperatures = filtered_df['Temperature'].dropna().tolist()
    else:
        print(f"Column 'Temperature' not found in file: {file_path}")
    print("Temperatures=", temperatures)
    return temperatures
  
def calculate_diff(start_time, end_time):
    start_datetime = pd.to_datetime(start_time)
    end_datetime = pd.to_datetime(end_time)
    diff_time = end_datetime - start_datetime
    # Extract total days, hours, and minutes
    days = diff_time.days
    total_seconds = diff_time.seconds
    hours, remainder = divmod(total_seconds, 3600)
    minutes = remainder // 60
    seconds = minutes // 60
    # Format the output to include days, hours, and minutes
    if days > 0:
        formatted_diff = f"{days} day(s), {int(hours)} hour(s),  {int(minutes)} minute(s) and {int(seconds)} second(s)"
    else:
        formatted_diff = f"{int(hours)} hour(s) and {int(minutes)} minute(s) and {int(seconds)} second(s)"
    return formatted_diff

def process_file(doc, files, start_datetime, end_datetime):
    index=1
    total_avg_time=0
    table = doc.add_table(rows=1, cols=4)
    table.autofit = True
    table.allow_autofit = True
    table.style = 'Table Grid'
    overall_min, overall_max = 0, 0
    update_style(table)
    table.cell(0, 0).text = 'Measuring Point'
    table.cell(0, 1).text = 'Average [°C]'
    table.cell(0, 2).text = 'Minimum [°C]'
    table.cell(0, 3).text = 'Maximum [°C]'
    min_max_avg_df = pd.DataFrame(columns=["Filename", "Min Temp", "Max Temp", "Avg Temp"])
    print("files=", files)
    for filename in files:
        temperatures = extract_temperatures_from_csv(filename, start_datetime, end_datetime)
        max_temp = max(temperatures)
        min_temp = min(temperatures)
        if index == 1:
            overall_max = max_temp
            overall_min = min_temp
        else:
            overall_min = min(overall_min, min_temp)
            overall_max = max(overall_max, max_temp)
        avg_temp = round((max_temp+min_temp)/2, 2)
        min_max_avg_df = pd.concat(
        [min_max_avg_df, pd.DataFrame({"Filename": [filename.name], "Min Temp": [min_temp], "Max Temp": [max_temp], "Avg Temp": [avg_temp]})],
        ignore_index=True,
         )
        total_avg_time += int(avg_temp)
        row_cells = table.add_row().cells
        row_cells[0].text = str(filename.name)
        row_cells[1].text = str(avg_temp) 
        row_cells[2].text = str(min_temp) 
        row_cells[3].text = str(max_temp)
        index+=1
    total_avg_time/=(index-1)
    return total_avg_time, overall_min, overall_max, min_max_avg_df

def get_min_max_average_graph(df):
  plt.figure(figsize=(20, 12))  # Set the figure size
  
  # Loop through each file and plot its min, avg, and max temperatures
  for i in range(len(df)):
    x = df['Filename'][i][0:7]
    temps = [df['Min Temp'][i], df['Avg Temp'][i], df['Max Temp'][i]]
    plt.plot(
        [x, x, x],  # Repeating the sensor name
        temps,  # Min, Avg, Max temps
        color='magenta',  # Vertical line color
        linestyle='-',  # Line style
        zorder=1
    )
    plt.scatter(x, df['Min Temp'][i], color='blue', label='Min Temp' if i == 0 else "", marker='s', zorder=2)
    plt.text(x, df['Min Temp'][i],  df['Min Temp'][i],
        fontsize=12, 
        color='blue', 
        ha='right', 
        va='bottom'
    )
    plt.scatter(x, df['Avg Temp'][i], color='green', label='Avg Temp' if i == 0 else "", marker='o', zorder=2)
    plt.text(x, df['Avg Temp'][i], df['Avg Temp'][i], 
        fontsize=12, 
        color='green', 
        ha='right', 
        va='bottom'
    )
    plt.scatter(x, df['Max Temp'][i], color='red', label='Max Temp' if i == 0 else "", marker='^', zorder=2)
    plt.text(x, df['Max Temp'][i], df['Max Temp'][i], 
        fontsize=12, 
        color='red', 
        ha='right', 
        va='bottom'
    )

  # Add labels and legend
  #ax.set_xticklabels(df["Filename"], fontsize=12, rotation=45, ha="right")
  plt.xticks(fontsize=15, rotation=45, ha='right')  # Rotate x-axis labels
  plt.xlabel('Sensor', fontsize=15)
  plt.ylabel('Temperature Degrees Centigrate', fontsize=18)
  #plt.title('Temperature Mapping - Individual Sensor Results')
  plt.legend(title='Legend', title_fontsize=15, fontsize=15)
  plt.grid(axis='y', linestyle='--', alpha=0.7)
  
  # Show the plot
  plt.tight_layout()
  plt.savefig("min_max_avg_plot.jpg")
  return "min_max_avg_plot.jpg"
  
def get_combined_graph(folder_path, start_datetime, end_datetime):
    # Step 1: Get all CSV files in the specified folder
    csv_files = folder_path
    num_files = len(csv_files)
    colors = plt.cm.viridis(np.linspace(0, 1, num_files))
    # Step 2: Create a new figure for plotting
    plt.figure(figsize=(16, 9))
    ax = plt.gca() 
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%m/%d/%Y\n%I:%M:%S %p'))
    # Initialize an empty DataFrame for combined data
    combined_data = pd.DataFrame()  
    # Step 3: Loop through each CSV file and filter data
     
    for idx, file in enumerate(csv_files):
          # Read the CSV file
          bytes_data = file.get_bytes()
          df = pd.read_csv(BytesIO(bytes_data)) 
          #df = pd.read_csv(BytesIO(bytes_data.decode('utf-8')))
          #df = pd.read_csv(StringIO(file.get_bytes().decode('utf-8')))
          #df = pd.read_csv(file)
          df['DateTime'] = pd.to_datetime(df['Date'] + ' ' + df['Time'])
          filtered_data = df[(df['DateTime'] >= start_datetime) & (df['DateTime'] <= end_datetime)]
          #filtered_data['SmoothedTemperature'] = filtered_data['Temperature'].rolling(window=5).mean()
          filtered_data['SmoothedTemperature'] = filtered_data['Temperature']
          # Check if filtered data is not empty before plotting
          if not filtered_data.empty:
              # Append the filtered data to the combined DataFrame
              combined_data = pd.concat([combined_data, filtered_data], ignore_index=True)
              # Plot the filtered data using the new DateTime column for x-axis
              plt.plot(filtered_data['DateTime'], 
                      filtered_data['SmoothedTemperature'], 
                      linestyle='-', color=colors[idx], label=file.name[0:7])
          else:
              print(f"No data in {file} for the given time period.")
    # Step 4: Add titles and labels
    plt.xlabel('Temperature (°C)')
    plt.ylabel('Date and Time')
    min_temp = filtered_data['Temperature'].min()
    max_temp = filtered_data['Temperature'].max()
    ax.set_yticks(np.arange(np.floor(min_temp), np.ceil(max_temp) + 1, 1))
    plt.xticks(rotation=45)  # Rotate x-axis labels for better visibility
    plt.legend()
    plt.grid()
    # Save and show the plot
    plt.tight_layout()
    plt.savefig("combined_data_plot.jpg")
    return "combined_data_plot.jpg", num_files

def low_high_calculator(folder_path, doc, temp, str_value, start_datetime, end_datetime):
    table = doc.add_table(rows=1, cols=3)
    update_style(table)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    table.cell(0, 0).text = 'Measuring Point(s) ' + str_value + ' Value'
    table.cell(0, 1).text = 'Value[°C]'
    table.cell(0, 2).text = 'Date / Time'
    for filename in folder_path:
          bytes_data = filename.get_bytes()
          df = pd.read_csv(BytesIO(bytes_data)) 
          df['DateTime'] = pd.to_datetime(df['Date'] + ' ' + df['Time'])
          filtered_df = df[(df['DateTime'] >= start_datetime) & (df['DateTime'] <= end_datetime) & (df['Temperature'] == temp)]
          if not filtered_df.empty:
              row = filtered_df.iloc[0]
              row_cells = table.add_row().cells
              row_cells[0].text = filename.name
              row_cells[1].text = str(row['Temperature'])
              row_cells[2].text = str(row['Date'] + " " + row['Time'])
            
def extract_temp_from_csv(file_path):
    temperatures = []
    # Read the CSV file into a DataFrame
    bytes_data = file_path.get_bytes()
    df = pd.read_csv(BytesIO(bytes_data)) 
    temperatures = df['Temperature'].dropna().tolist()
    return temperatures

def process_all_files(folder_path, start_datetime, end_datetime):
    # Convert start and end date/time strings to datetime objects
    start_datetime = pd.to_datetime(start_datetime)
    end_datetime = pd.to_datetime(end_datetime)
    
    all_temperatures = []
    
    for file_obj in folder_path:
        # Check if file_obj is a StreamingMedia object and convert it
        if hasattr(file_obj, 'get_bytes'):
            file_data = BytesIO(file_obj.get_bytes())
        else:
            raise ValueError("Unsupported file object type.")
        
        # Load the CSV file into a DataFrame
        df = pd.read_csv(file_data)
        
        # Combine 'Date' and 'Time' into a single 'datetime' column if needed
        if 'datetime' not in df.columns:
            df['datetime'] = pd.to_datetime(df['Date'] + ' ' + df['Time'])
        
        # Filter rows within the date-time range
        filtered_df = df[(df['datetime'] >= start_datetime) & (df['datetime'] <= end_datetime)]
        
        # Collect the temperatures
        all_temperatures.extend(filtered_df['Temperature'].tolist())
    if all_temperatures:
        max_temp = max(all_temperatures)
        min_temp = min(all_temperatures)
        avg_temp = round(((max_temp+min_temp)/2), 2)
        median_temp=statistics.median(all_temperatures)
        return max_temp, min_temp, avg_temp,median_temp
    else:
        print(f"No temperature data found between {start_datetime} and {end_datetime} across all files.")

def generate_graph_for_time_range(filtered_df, image_path):
    if not filtered_df.empty:
        plt.figure(figsize=(10, 6))
        ax = plt.gca() 
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%m/%d/%Y\n%I:%M:%S %p'))
        filtered_df.loc[:, 'DateTime'] = pd.to_datetime(filtered_df['Date'] + ' ' + filtered_df['Time'])
        filtered_df['SmoothedTemperature'] = filtered_df['Temperature'].rolling(window=5).mean()
    
        #plt.plot(filtered_df.loc[:, 'DateTime'], filtered_df['Temperature'], linestyle='-', color='b')
        plt.xlabel("Date and Time")
        plt.ylabel("Temperature (°C)")
        min_temp = filtered_df['Temperature'].min()
        max_temp = filtered_df['Temperature'].max()
        ax.set_yticks(np.arange(np.floor(min_temp), np.ceil(max_temp) + 1, 1))
        # Plot smoothed line, dropping NaN values that might arise from the rolling window
        plt.plot(filtered_df['DateTime'], filtered_df['SmoothedTemperature'], linestyle='-', color='orange')
        ax.axhspan(14, 16, facecolor='blue', alpha=0.1)
        ax.axhspan(16, 20, facecolor='green', alpha=0.1)
        ax.axhspan(20, 26, facecolor='red', alpha=0.1)
        #plt.grid(True)
        #plt.tight_layout(pad=1.5, w_pad=3.5, h_pad=1.0)
        img_stream = BytesIO()
        plt.savefig(img_stream, format='jpg')
        plt.close()
        img_stream.seek(0)  # Move to the start of the BytesIO object
        return img_stream
    else:
        print("No data available for the specified time range.")
        return ""
      
def mean_kinetic_temperature(temperatures, delta_h=83144, r=8.314):
    # Convert temperatures to Kelvin
    temps_in_kelvin = [temp + 273.15 for temp in temperatures]
    # Calculate the exponentials for each temperature
    exponentials = [math.exp(-delta_h / (r * temp)) for temp in temps_in_kelvin]
    # Calculate MKT using the formula
    mkt_kelvin = (-delta_h / r) / math.log(sum(exponentials) / len(exponentials))
    # Convert MKT back to °C
    mkt_celsius = mkt_kelvin - 273.15
    return round(mkt_celsius, 2)

def read_and_filter_data(doc, folder_path, start_datetime, end_datetime):
    counter=3.1
    value=0.1
    for filename in folder_path:
        if counter + value >= 5.0:
          counter = 4.1
          value = 0.1/10
          counter += value
        else:
          counter += value
        if value == 0.1:
          counter = round(counter, 1)
        elif value == 0.01:
          counter = round(counter, 2)
        heading = doc.add_heading(f'{str(counter)}: {filename.name}', level=2)
        update_heading_style(heading)
        bytes_data = filename.get_bytes()
        df = pd.read_csv(BytesIO(bytes_data)) 
        df['DateTime'] = pd.to_datetime(df['Date'] + ' ' + df['Time'])
        filtered_df = df[(df['DateTime'] >= start_datetime) & (df['DateTime'] <= end_datetime)]
        file_name = filename.name 
        base = os.path.splitext(file_name)[0]
        new_filename = base + ".jpg"
        new_file_path = new_filename
        image_path = generate_graph_for_time_range(filtered_df, new_file_path)
        doc.add_picture(image_path, width=Inches(6.0)) 
        doc.add_heading('', level=2)
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        table.autofit = True
        table.allow_autofit = True
        update_style(table)
        table.cell(0, 0).text = 'Statistics'
        table.cell(0, 1).text = 'Value'
        table.cell(1, 0).text = 'Range'
        table.cell(1, 1).text = str(start_datetime) + '-' + str(end_datetime)
        temperatures = extract_temperatures_from_csv(filename, start_datetime, end_datetime)
        max_temp = max(temperatures)
        min_temp = min(temperatures)
        avg_temp = round((max_temp+min_temp)/2, 2)
        table.cell(2, 0).text = 'Average'
        table.cell(2, 1).text = str(avg_temp) + ' °C'
        table.cell(3, 0).text = 'Min/Max Thresholds'
        table.cell(3, 1).text = str(min_temp) + ' / ' + str(max_temp) +  ' °C'
        calculated_mkt = mean_kinetic_temperature(temperatures)
        table.cell(4, 0).text = 'MKT'
        table.cell(4, 1).text = str(calculated_mkt) +  ' °C'
        variance = round(statistics.variance(temperatures), 2)
        table.cell(5, 0).text = 'Variance'
        table.cell(5, 1).text = f"{variance:.2f}"
        std_deviation = statistics.variance(temperatures)
        table.cell(6, 0).text = 'Std. Deviation'
        table.cell(6, 1).text = f"{std_deviation:.4f}"
        table.cell(7, 0).text = 'Time: Within Min/Max'
        diff_time = calculate_diff(start_datetime, end_datetime)
        table.cell(7, 1).text = diff_time
def add_equipment_table(doc, equipment_type, document_number, creation_date):
  table = doc.add_table(rows=2, cols=2)
  table.style = 'Table Grid'
  table.autofit = True
  table.allow_autofit = True
  update_style(table)
  table.cell(0, 0).text=f'Equipment: {equipment_type}'
  table.cell(1, 0).text=f'Document: {document_number}'
  table.cell(1, 1).text=f'Date: {creation_date}'
  
def add_table_of_contents(doc):
  title = doc.add_heading('Table of Contents ')
  update_heading_style(title)
  table = doc.add_table(rows=17, cols=2)
  update_style(table)
  table.style = 'Table Grid' 
  table.allow_autofit = True
# Remove table borders
  table.cell(0, 0).text = "SECTION 1: EXECUTIVE SUMMARY"
  table.cell(0, 1).text = "4"
  table.cell(1, 0).text = "SECTION 2: APPROVAL SIGNATURES"
  table.cell(1, 1).text = "5"
  table.cell(2, 0).text = "SECTION 3: REVISION HISTORY"
  table.cell(2, 1).text = "6"
  table.cell(3, 0).text = "SECTION 4: SIGNATURE CONTROL FORM"
  table.cell(3, 1).text = "6"
  cell = table.cell(4, 0)
  run = cell.paragraphs[0].add_run('SECTION 5: PROGRAM FORMAT')
  run = cell.paragraphs[0].add_run('\n5.1: Objective')
  run.italic=True
  run = cell.paragraphs[0].add_run('\n5.1: Objective')
  run.italic=True
  run = cell.paragraphs[0].add_run('\n5.2: Scope')
  run.italic=True
  run = cell.paragraphs[0].add_run('\n5.3: Rationale')
  run.italic=True
  run = cell.paragraphs[0].add_run('\n5.4: Responsibility')
  run.italic=True
  run = cell.paragraphs[0].add_run('\n5.5: Abbreviation/Glossary')
  run.italic=True
  run = cell.paragraphs[0].add_run('\n5.6: Deviations/Deficiencies')
  run.italic=True
  table.cell(4, 1).text = "7"
  cell = table.cell(5, 0)
  run = cell.paragraphs[0].add_run('SECTION 6: CALIBRATION')
  run = cell.paragraphs[0].add_run('\n6.1: Critical Instrumentation')
  run.italic=True
  run = cell.paragraphs[0].add_run('\n6.2: Testing Instrumentation')
  run.italic=True
  table.cell(5, 1).text = "10"
  table.cell(6, 0).text = "SECTION 7: STANDARD OPERATING PROCEDURES (SOP) LIST"
  table.cell(6, 1).text = "11"
  table.cell(7, 0).text = "SECTION 8: CRITICAL OPERATING PARAMETERS"
  table.cell(7, 1).text = "11"
  cell = table.cell(8, 0)
  run = cell.paragraphs[0].add_run('SECTION 9: TEST FUNCTIONS')
  run = cell.paragraphs[0].add_run('\n9.1: Test Function No. 1: Control Panel Verification')
  run.italic=True
  run = cell.paragraphs[0].add_run('\n9.2: Test Function No. 2: Major Components Operating Parameters Verification')
  run.italic=True
  run = cell.paragraphs[0].add_run('\n9.3: Test Function No. 3: Empty Trailer Temperature Distribution Verification')
  run.italic=True
  run = cell.paragraphs[0].add_run('\n9.4: Test Function No. 4: Trailer Recovery Verification')
  run.italic=True
  table.cell(8, 1).text = "11"
  table.cell(9, 0).text = "SECTION 10: ANALYSIS AND CONCLUSIONS"
  table.cell(9, 1).text = "15"
  table.cell(10, 0).text = "APPENDIX A:DRAWINGS"
  table.cell(10, 1).text = "16"
  table.cell(11, 0).text = "APPENDIX B:TEST RESULTS TABLES AND GRAPHS"
  table.cell(11, 1).text = "17"
  table.cell(12, 0).text = "APPENDIX C:RAW DATA SHEETS"
  table.cell(12, 1).text = "18"
  table.cell(13, 0).text = "APPENDIX D:CRITICAL INSTRUMENT CALIBRATION REPORTS"
  table.cell(13, 1).text = "19"
  table.cell(14, 0).text = "APPENDIX E:TESTING INSTRUMENT CALIBRATION REPORTS"
  table.cell(14, 1).text = "20"
  table.cell(15, 0).text = "APPENDIX F:REFERENCE/RELATED DOCUMENTS"
  table.cell(15, 1).text = "21"
  table.cell(16, 0).text = "APPENDIX G:DEVIATION/DEFICIENCY REPORTS"
  table.cell(16, 1).text = "22"
  
def create_document(files, start_datetime, end_datetime, start_input, set_point, company_name, author_name, app_name, trailer_no, season, protocol_number, document_number, make,model_number,vin_number, equipment_number, operating_conditions, creation_date, equipment_type, revision_date):
    doc = Document()
    #logo_width = Inches(4.5)
    #image=doc.add_paragraph()
    #run=image.add_run()
    #run.add_picture("skyline logo.png",width=logo_width)
    if start_input == 1:
        title = doc.add_heading('6-Hour Mapping-Empty Trailer', level=1)
    elif start_input == 2:
        title = doc.add_heading('6-Hour Mapping-Loaded Trailer', level=1)
    elif start_input == 3:
        title = doc.add_heading('1-Hour Temperature-Control Power Failure-Loaded Trailer', level=1)
    elif start_input == 4:
        title = doc.add_heading('1-Minute Open Door-Loaded Trailer', level=1)
    else:
        title = doc.add_heading('2-Hour Field Shipment Test-Loaded Trailer', level=1)
    update_heading_style(title)
    #title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #second_title = doc.add_heading(company_name + ' ' +  str(set_point) + ' °C', level=2)
    #second_title.italic = True
    #second_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #update_heading_style(second_title)

    title = doc.add_heading('Operational Qualification (OQ): Refrigerated Trailer', level=1)
    update_heading_style(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=11, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    cell = table.cell(0, 0)
    run = cell.paragraphs[0].add_run('Protocol Number')
    run.bold = True
    cell = table.cell(0,1)
    run = cell.paragraphs[0].add_run(protocol_number)
    cell = table.cell(1, 0)
    run = cell.paragraphs[0].add_run('Document Number')
    run.bold = True
    cell = table.cell(1,1)
    run = cell.paragraphs[0].add_run(document_number)
    cell = table.cell(2, 0)
    run = cell.paragraphs[0].add_run('Date Prepared')
    run.bold = True
    cell = table.cell(2,1)
    run = cell.paragraphs[0].add_run(str(creation_date))
    cell = table.cell(3, 0)
    run = cell.paragraphs[0].add_run('Equipment Type')
    run.bold = True
    cell = table.cell(3,1)
    run = cell.paragraphs[0].add_run("Refrigerated Trailer")
    cell = table.cell(4, 0)
    run = cell.paragraphs[0].add_run('Make')
    run.bold = True
    cell = table.cell(4,1)
    run = cell.paragraphs[0].add_run(make)
    cell = table.cell(5, 0)
    run = cell.paragraphs[0].add_run('Model Number')
    run.bold = True
    cell = table.cell(5,1)
    run = cell.paragraphs[0].add_run(model_number)
    cell = table.cell(6, 0)
    run = cell.paragraphs[0].add_run('VIN Number')
    run.bold = True
    cell = table.cell(6,1)
    run = cell.paragraphs[0].add_run(vin_number)
    cell = table.cell(7, 0)
    run = cell.paragraphs[0].add_run('Equipment Id No')
    run.bold = True
    cell = table.cell(7,1)
    run = cell.paragraphs[0].add_run(equipment_number)
    cell = table.cell(8, 0)
    run = cell.paragraphs[0].add_run('Operating Conditions')
    run.bold = True
    cell = table.cell(8,1)
    run = cell.paragraphs[0].add_run(str(operating_conditions))
    cell = table.cell(9, 0)
    run = cell.paragraphs[0].add_run('Revision Date')
    run.bold = True
    cell = table.cell(9,1)
    run = cell.paragraphs[0].add_run(str(revision_date))
    cell = table.cell(10, 0)
    run = cell.paragraphs[0].add_run('Owner')
    run.bold = True
    cell = table.cell(10,1)
    run = cell.paragraphs[0].add_run('Skyline Cargo \n7027 Fir Tree Drive \nMississauga, ON \nL5S 1J7 ')
    update_style(table)
    add_table_of_contents(doc)
    add_equipment_table(doc, equipment_number, document_number, creation_date)
    title = doc.add_heading('SECTION 1: EXECUTIVE SUMMARY', level=1)
    update_heading_style(title)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid' 
    table.allow_autofit = True
    cell = table.cell(0, 0)
    run = cell.paragraphs[0].add_run('Make: ')
    run.bold = True
    cell = table.cell(0, 1)
    run = cell.paragraphs[0].add_run('Model: ')
    run.bold = True
    cell = table.cell(0, 2)
    run = cell.paragraphs[0].add_run('Serial: ')
    run.bold = True
    cell = table.cell(1, 0)
    run = cell.paragraphs[0].add_run('Client:')
    run.bold = True
    run = cell.paragraphs[0].add_run('Skyline cargo:')
    cell = table.cell(1, 1)
    run = cell.paragraphs[0].add_run('Client Equipment ID: ')
    run.bold = True
    cell = table.cell(2, 0)
    run = cell.paragraphs[0].add_run('#1')
    cell = table.cell(2, 1).merge(table.cell(2, 2))
    run = cell.paragraphs[0].add_run('All features on the control panel operated according to design specifications.\n\nRESULTS:\n☐ PASS | ☐ FAIL\n Deviation Report No: _____________________     Deficiency Report No: _____________________')
    cell = table.cell(3, 0)
    run = cell.paragraphs[0].add_run('#2')
    cell = table.cell(3, 1).merge(table.cell(3, 2))
    run = cell.paragraphs[0].add_run('The unit’s major components operated according to design specifications.\n\nRESULTS:\n☐ PASS | ☐ FAIL\n Deviation Report No: _____________________     Deficiency Report No: _____________________')
    cell = table.cell(4, 0)
    run = cell.paragraphs[0].add_run('#3')
    cell = table.cell(4, 1).merge(table.cell(4, 2))
    run = cell.paragraphs[0].add_run('For the six (6) hour mapping performed at the empty trailer set-up conditions (2°C to 8°C), The temperature readings recorded by each of the sensors located in the trailer remained within (2°C to 8°C).\n\nRESULTS:\n☐ PASS | ☐ FAIL\n Deviation Report No: _____________________     Deficiency Report No: _____________________')
    cell = table.cell(5, 0)
    run = cell.paragraphs[0].add_run('#4')
    cell = table.cell(5, 1).merge(table.cell(5, 2))
    run = cell.paragraphs[0].add_run('For the six (6) hour mapping performed at the empty trailer set-up conditions (15°C to 25°C), The temperature readings recorded by each of the sensors located in the trailer remained within (15°C to 25°C). \n\nRESULTS:\n☐ PASS | ☐ FAIL\n Deviation Report No: _____________________     Deficiency Report No: _____________________')
    doc.add_paragraph()
    add_equipment_table(doc, equipment_number, document_number, creation_date)
    doc.add_paragraph()
    title = doc.add_heading('SECTION 2: APPROVAL SIGNATURES', level=1)
    update_heading_style(title)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid' 
    table.allow_autofit = True
    cell = table.cell(0, 0)
    run = cell.paragraphs[0].add_run('Document Name: ')
    run.bold = True
    cell = table.cell(0, 1)
    run = cell.paragraphs[0].add_run('Document No.: ')
    run.bold = True
    cell = table.cell(1, 0)
    run = cell.paragraphs[0].add_run('Make:')
    run.bold = True
    cell = table.cell(1, 1)
    run = cell.paragraphs[0].add_run('Model: ')
    run.bold = True
    cell = table.cell(1, 2)
    run = cell.paragraphs[0].add_run('Serial: ')
    run.bold = True
    cell = table.cell(2, 0)
    run = cell.paragraphs[0].add_run('Client: Skyline Cargo')
    cell = table.cell(2, 1)
    run = cell.paragraphs[0].add_run('Contact: Paul Budwal')
    cell = table.cell(3, 0)
    run = cell.paragraphs[0].add_run('Post-Execution Protocol Review:\n')
    run.bold = True
    cell.paragraphs[0].add_run('These signatures indicate that the protocol for the equipment designated above has been completed. The results have been reviewed by the client and found to satisfy all protocol requirements. Approval for the continued use of this equipment is hereby granted. ')
    cell = table.cell(4, 0)
    run = cell.paragraphs[0].add_run('NAME, TITLE, DEPARTMENT')
    run.bold = True
    cell = table.cell(4, 1)
    run = cell.paragraphs[0].add_run('SIGNATURE')
    run.bold = True
    cell = table.cell(4, 2)
    run = cell.paragraphs[0].add_run('DATE')
    run.bold = True
    cell = table.cell(5, 0)
    run = cell.paragraphs[0].add_run('Paul Budwal\nTitle - QA   Skyline Cargo')
    cell = table.cell(6, 0)
    run = cell.paragraphs[0].add_run('Kam Budwal\nTitle - CEO   Skyline Cargo')
    doc.add_paragraph()
    add_equipment_table(doc, equipment_number, document_number, creation_date)
    doc.add_paragraph()
    title = doc.add_heading('SECTION 4: SIGNATURE CONTROL FORM', level=1)
    update_heading_style(title)
    table = doc.add_table(rows=8, cols=5)
    table.style = 'Table Grid' 
    table.allow_autofit = True
    cell = table.cell(0, 0).merge(table.cell(0, 4))
    run = cell.paragraphs[0].add_run('All personnel who are involved in the execution and review of this protocols study results must enter samples of their signature and initials in the table below.')
    cell = table.cell(1, 0)
    run = cell.paragraphs[0].add_run('NAME')
    run.bold = True
    cell = table.cell(1, 1)
    run = cell.paragraphs[0].add_run('TITLE')
    run.bold = True
    cell = table.cell(1, 2)
    run = cell.paragraphs[0].add_run('SIGNATURE')
    run.bold = True
    cell = table.cell(1, 3)
    run = cell.paragraphs[0].add_run('INITIALS')
    run.bold = True
    cell = table.cell(1, 4)
    run = cell.paragraphs[0].add_run('DATE')
    run.bold = True
    for row in range(2, 8):
      for column in range(0, 5):
        cell = table.cell(row, column)
        run = cell.paragraphs[0].add_run('')
    doc.add_paragraph()
    add_equipment_table(doc, equipment_number, document_number, creation_date)
    doc.add_paragraph()
    title = doc.add_heading('SECTION 5: PROGRAM FORMAT', level=1)
    update_heading_style(title)
    table = doc.add_table(rows=4, cols=1)
    table.style = 'Table Grid' 
    table.allow_autofit = True
    cell = table.cell(0, 0)
    run = cell.paragraphs[0].add_run('5.1: Objective')
    run.bold = True
    cell.paragraphs[0].add_run('\nThe objective of this protocol is to define and document the evidence needed to verify that the HYTR THE Trailer is operating according to HYTR and Skyline Cargo specifications, permitting operation as per design specifications as well as to cGMP.')
    cell = table.cell(2, 0)
    run = cell.paragraphs[0].add_run('5.2: Scope')
    run.bold = True
    cell.paragraphs[0].add_run('\nThe scope of this Operational Qualification study is limited to the equipment, instruments and components described in this protocol at the location stated in this protocol. ')
    cell = table.cell(3, 0)
    run = cell.paragraphs[0].add_run('5.3: Rationale')
    run.bold = True
    cell.paragraphs[0].add_run('\nThe Operational Qualification Protocol is designed to allow technical analysis of all applicable operational functions. The proper operation of the Trailer will be established based on the following characteristics and rationale:')
    run = cell.paragraphs[0].add_run('\nTest Function No. 1:')
    run.bold = True
    cell.paragraphs[0].add_run('The operation of the control panel will be verified to demonstrate that it operates according to design specifications.')
    run = cell.paragraphs[0].add_run('\nTest Function No. 2:')
    run.bold = True
    cell.paragraphs[0].add_run('The operation of the major components will be verified by testing the various inputs and outputs to and from those components.')
    run = cell.paragraphs[0].add_run('\nTest Function No. 3:')
    run.bold = True
    cell.paragraphs[0].add_run('The temperature control and distribution within the Trailer will be demonstrated by conducting a four-hour mapping under empty Trailer conditions. The mapping will be conducted for the operating ranges of (2°C to 8°C) and (15°C to 25°C) at which the Trailer will be used. The reading interval is set at 1 minutes for the mapping study. To demonstrate that the Trailer temperature can remain within the established specifications for prolonged time periods, the temperatures recorded by each of the sensors located in the Trailer must remain within (2°C to 8°C) or (15°C to 25°C).')
    doc.add_paragraph()
    title = doc.add_heading('SECTION 5: PROGRAM FORMAT(Continued)', level=1)
    update_heading_style(title)
    table = doc.add_table(rows=3, cols=1)
    table.style = 'Table Grid' 
    table.allow_autofit = True
    cell = table.cell(0, 0)
    run = cell.paragraphs[0].add_run('5.4 : Responsibility ')
    run.bold = True
    cell.paragraphs[0].add_run('The following responsibilities were assigned for the operation qualification of the chamber:\n') 
    run = cell.paragraphs[0].add_run('Responsibility                                QA                                  Management\n') 
    run.bold = True
    cell.paragraphs[0].add_run('\n\nWrite the protocol')
    cell.paragraphs[0].add_run('\nReview and approve the protocol')
    cell.paragraphs[0].add_run('\nCalibrate the critical instruments')
    cell.paragraphs[0].add_run('\nCalibrate and supply the testing instruments')
    cell.paragraphs[0].add_run('\nEnsure that instruments are in calibrated state')
    cell.paragraphs[0].add_run('\nProvide the functional specifications')
    cell.paragraphs[0].add_run('\nConduct the testing as described in the protocol')
    cell.paragraphs[0].add_run('\nCompile and analyze the test data')
    cell.paragraphs[0].add_run('\nIssue the final validation report')
    cell.paragraphs[0].add_run('\nReview and approve the final report')
    cell.paragraphs[0].add_run('\Review and approve the deviation/deficiency report(s)')
    cell = table.cell(1, 0)
    run = cell.paragraphs[0].add_run('5.5: Abbreviation Glossary ')
    run.bold = True
    cell.paragraphs[0].add_run('\nAbbreviation            Description')
    cell.paragraphs[0].add_run('\nA                        Amperes')
    cell.paragraphs[0].add_run('\ncGMP                    Current Good Manufacturing Practice')
    cell.paragraphs[0].add_run('\nHz                       H')
    cell.paragraphs[0].add_run('\nHz                       H')
    cell.paragraphs[0].add_run('\nNAP                     Not Applicable')
    cell.paragraphs[0].add_run('\nNAV                     Not Available')
    cell.paragraphs[0].add_run('\nNIST                    National Institute of Standards and Technology (USA)')
    cell.paragraphs[0].add_run('\nNRC                     National Research Council (Canada)')
    cell.paragraphs[0].add_run('\nph                        Phase')
    cell.paragraphs[0].add_run('\nRH                       Relative Humidity')
    cell.paragraphs[0].add_run('\nRTD                     Resistance Temperature Device')
    cell.paragraphs[0].add_run('\nSOP                      Standard Operating Procedure')
    cell.paragraphs[0].add_run('\nTBD                     To Be Determined')
    cell.paragraphs[0].add_run('\nTC                       Thermocouple')
    cell.paragraphs[0].add_run('\nVAC                     Volts Alternating Current')
    cell.paragraphs[0].add_run('\nVDC                     Volts Direct Current')
    doc.add_paragraph()
    add_equipment_table(doc, equipment_number, document_number, creation_date)
    doc.add_paragraph()
    title = doc.add_heading('SECTION 5: PROGRAM FORMAT(Continued)', level=1)
    update_heading_style(title)
    table = doc.add_table(rows=2, cols=1)
    table.style = 'Table Grid' 
    table.allow_autofit = True
    cell = table.cell(0, 0)
    run = cell.paragraphs[0].add_run('5.6 : Deviations / Deficiencies  ')
    run.bold=True
    cell.paragraphs[0].add_run('\nIt is possible that during the execution of the validation protocol, that the pre-established test method may not be feasible. The changes made during execution are documented in a deviation report. The deviation report is used to document a deviation from an approved protocol. A deviation from an approved protocol is defined as, for example: a test methodology or acceptance criteria that was changed during execution but has no impact on the validity of the test. The deviation report describes the deviation, and the rationale is explained. ')
    cell.paragraphs[0].add_run('\n\nA Deficiency is defined as a non-conformity identified during a protocol execution and that prevents from meeting the pre-established and approved acceptance criteria. A Deficiency is a non-planned event, which appears during a protocol execution. The deficiency report is used to identify any failure of the equipment, specifications or test results that do not meet the acceptance criteria of an approved validation protocol. As a rule, the deficiency is described by the person who performed validation of the equipment. ')
    cell.paragraphs[0].add_run('\n\nThe impact on the system, the magnitude of the deficiency and the need for corrective action are evaluated by the most relevant person responsible concerning the actual deficiency. ')
    cell.paragraphs[0].add_run('\n\nA critical deficiency is defined as critical regarding the quality of the product or the function of a critical piece of equipment or system. A critical deficiency prevents a protocol/final report from being approved and must be resolved before its approval.')
    cell.paragraphs[0].add_run('\n\nIn case of a non-critical deficiency, the final report is approved conditionally, and an action plan is put in place in order to make sure that the corrective measures will be completed and documented. A non-critical deficiency does not affect the quality of the product or the function of a critical piece of equipment or system')
    doc.add_paragraph()
    add_equipment_table(doc, equipment_number, document_number, creation_date)
    doc.add_paragraph()
    title = doc.add_heading('SECTION 6: CALIBRATION', level=1)
    update_heading_style(title)
    table = doc.add_table(rows=24, cols=2)
    table.style = 'Table Grid' 
    table.allow_autofit = True
    cell = table.cell(0, 0).merge(table.cell(0, 1))
    run = cell.paragraphs[0].add_run('SECTION 6.1 : Critical Instrumentation')
    run.bold=True
    cell.paragraphs[0].add_run('\nThe following testing instrumentation were used during the execution and must have been calibrated using equipment that was traceable to national standards (NIST, NRC), within the last 12 months prior to execution. The thermocouples (temperature sensors) must have been calibrated prior to and verified after the execution of this qualification.')
    cell = table.cell(1, 0)
    run = cell.paragraphs[0].add_run('Instrument/Sensor Description')
    run.bold = True
    cell = table.cell(1, 1)
    run = cell.paragraphs[0].add_run("Calibration Date")
    run.bold = True
    cell = table.cell(2, 0)
    cell.paragraphs[0].add_run('Temperature Controller')
    cell = table.cell(2, 1)
    cell.paragraphs[0].add_run('')
    cell = table.cell(3, 0)
    cell.paragraphs[0].add_run('Monitoring System Temperature Sensor ')
    cell = table.cell(3, 1)
    cell.paragraphs[0].add_run('')
    cell = table.cell(4, 0)
    cell.paragraphs[0].add_run('Calibration certificates and reports relating to the calibration of the unit’s instrumentation will be annexed in Appendix D: Critical Instrument Calibration Reports')
    cell = table.cell(5, 0).merge(table.cell(5,1))
    run = cell.paragraphs[0].add_run('SECTION 6.2 : Testing Instrumentation ')
    run.bold=True
    cell.paragraphs[0].add_run('The following testing instrumentation were used during the execution and must have been calibrated using equipment that was traceable to national standards (NIST, NRC), within the last 12 months prior to execution. The thermocouples (temperature sensors) must have been calibrated prior to and verified after the execution of this qualification.')
    cell = table.cell(6, 0)
    run = cell.paragraphs[0].add_run('Instrument/Sensor Description')
    run.bold = True
    cell = table.cell(6, 1)
    run = cell.paragraphs[0].add_run("Calibration Date")
    run.bold = True
    for i in range(7, 23):
      for j in range(0, 2):
        cell = table.cell(i, j)
        cell.paragraphs[0].add_run('')
    cell = table.cell(23, 0).merge(table.cell(23, 1))
    run = cell.paragraphs[0].add_run('Calibration certificates that relate to the calibration of the testing instrumentation will be annexed in Appendix E: Testing Instrument Calibration Reports. ')
    doc.add_paragraph()
    add_equipment_table(doc, equipment_number, document_number, creation_date)
    doc.add_paragraph()
    title = doc.add_heading('SECTION 7: STANDARD OPERATING PROCEDURES (SOP) LIST', level=1)
    update_heading_style(title)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid' 
    table.allow_autofit = True
    cell = table.cell(0, 0).merge(table.cell(0,3))
    cell.paragraphs[0].add_run('The following table lists all standard operating procedures that are directly applicable to the operation of the unit under study. All procedures must be completed prior to the approval of the operational qualification report. SOPs should include – Operation, Cleaning, Maintenance and Calibration of various critical instruments/devices. Respective personnel should be trained according to these SOP’s')
    cell = table.cell(1, 0)
    run = cell.paragraphs[0].add_run('Document Number')
    run.bold = True
    cell = table.cell(1, 1)
    run = cell.paragraphs[0].add_run('Revision Number')
    run.bold = True
    cell = table.cell(1, 2)
    run = cell.paragraphs[0].add_run('Title')
    run.bold = True
    cell = table.cell(1, 3)
    run = cell.paragraphs[0].add_run('Status')
    run.bold = True
    cell = table.cell(2, 0).merge(table.cell(2,3))
    run = cell.paragraphs[0].add_run('Preventative Maintenance Policy')
    cell = table.cell(3, 0)
    cell.paragraphs[0].add_run('STATUS LEGEND: \n')
    cell.paragraphs[0].add_run("IP =	In-process        TBW = To Be Written\n")
    cell.paragraphs[0].add_run("R =	Required            TBR = To Be Revised\n") 
    cell.paragraphs[0].add_run("C =	Completed\n")
    doc.add_paragraph()
    title = doc.add_heading('SECTION 8: CRITICAL OPERATING PARAMETERS', level=1)
    update_heading_style(title)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid' 
    table.allow_autofit = True
    cell = table.cell(0, 0).merge(table.cell(0,2))
    cell.paragraphs[0].add_run('The following parameters may have an impact on the proper operation of the Trailer. ')
    cell = table.cell(1, 0)
    run = cell.paragraphs[0].add_run('Parameter')
    run.bold = True
    cell = table.cell(1, 1)
    run = cell.paragraphs[0].add_run('Operating Range/Value')
    run.bold = True
    cell = table.cell(1, 2)
    run = cell.paragraphs[0].add_run('Range/Value to be covered during this study')
    run.bold = True
    cell = table.cell(2, 0)
    run = cell.paragraphs[0].add_run('Tailer Load Configuration')
    cell = table.cell(2, 1)
    run = cell.paragraphs[0].add_run('Empty to Fully Loaded')
    cell = table.cell(2, 2)
    run = cell.paragraphs[0].add_run('Empty')
    cell = table.cell(3, 0)
    cell.paragraphs[0].add_run('Temperature Setpoint')
    cell = table.cell(3, 1)
    cell.paragraphs[0].add_run('NAV')
    cell = table.cell(3, 2)
    cell.paragraphs[0].add_run('(2°C to 8°C) & (15°C to 25°C)')
    doc.add_paragraph()
    title = doc.add_heading('SECTION 9: TEST FUNCTIONS', level=1)
    update_heading_style(title)
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid' 
    table.allow_autofit = True
    cell = table.cell(0, 0)
    cell.paragraphs[0].add_run('Test Function No. 1: Control Panel Verification')
    cell.paragraphs[0].add_run('\nTest Function No. 2: Major Components Operating Parameters Verification')
    cell.paragraphs[0].add_run('\nTest Function No. 3: Empty Trailer Temperature Distribution')
    doc.add_paragraph()
    title = doc.add_heading('SECTION 9: TEST FUNCTIONS (Continued)', level=1)
    update_heading_style(title)
    table = doc.add_table(rows=19, cols=2)
    table.style = 'Table Grid' 
    table.allow_autofit = True
    cell = table.cell(0, 0).merge(table.cell(0,1))
    run = cell.paragraphs[0].add_run('9.1 Test Function No. 1: Control Panel Verification')
    run.bold = True
    cell = table.cell(1, 0).merge(table.cell(1,1))
    run = cell.paragraphs[0].add_run('Objective: ')
    run.bold = True
    cell.paragraphs[0].add_run('To demonstrate that the control panel features operate as per design specifications.')
    cell = table.cell(2, 0).merge(table.cell(2,1))
    run = cell.paragraphs[0].add_run('Acceptance Criteria: ')
    run.bold = True
    cell.paragraphs[0].add_run('All features of the control panel must operate as described in the Methodology Section.')
    cell = table.cell(3, 0).merge(table.cell(3,1))
    run = cell.paragraphs[0].add_run('Testing Equipment / Instruments: ')
    run.bold = True
    cell.paragraphs[0].add_run('None Required')
    cell = table.cell(4, 0).merge(table.cell(4,1))
    run = cell.paragraphs[0].add_run('Methodology: ')
    run.bold = True
    cell.paragraphs[0].add_run('None Required')
    cell = table.cell(5, 0).merge(table.cell(5,1))
    run = cell.paragraphs[0].add_run('Control Panel ')
    run.bold = True
    cell = table.cell(6, 0)
    run = cell.paragraphs[0].add_run('Test ')
    run.bold = True
    cell = table.cell(6, 1)
    run = cell.paragraphs[0].add_run('Conform Y/N ')
    run.bold = True
    cell = table.cell(7, 0)
    run = cell.paragraphs[0].add_run('Main temperature display -  ')
    run.bold = True
    cell.paragraphs[0].add_run('during normal operation, shows cabinet temperature in degrees Celsius, when the compressor is running, type of fuel, heating/cooling, and mode')
    cell = table.cell(7, 1)
    cell.paragraphs[0].add_run(' ')
    cell = table.cell(8, 0)
    cell.paragraphs[0].add_run('– Increases the value of the selected setting')
    cell = table.cell(8, 1)
    cell.paragraphs[0].add_run(' ')
    cell = table.cell(9, 0)
    cell.paragraphs[0].add_run('- Saves a change to the selected value.')
    cell = table.cell(9, 1)
    cell.paragraphs[0].add_run(' ')
    cell = table.cell(10, 0)
    cell.paragraphs[0].add_run('- Decreases the value of the selected setting.')
    cell = table.cell(10, 1)
    cell.paragraphs[0].add_run(' ')
    cell = table.cell(11, 0)
    cell.paragraphs[0].add_run('- Menu – switches between the various mode screens')
    cell = table.cell(11, 1)
    cell.paragraphs[0].add_run(' ')
    cell = table.cell(12, 0)
    cell.paragraphs[0].add_run('- Defrost – Initiates a defrost cycle')
    cell = table.cell(12, 1)
    cell.paragraphs[0].add_run(' ')
    cell = table.cell(13, 0)
    run = cell.paragraphs[0].add_run('Start Stop  Continuous')
    run.bold = True
    cell.paragraphs[0].add_run('– switches between controlled operation and continuous cooling')
    cell = table.cell(13, 1)
    cell.paragraphs[0].add_run(' ')
    cell = table.cell(14, 0)
    cell.paragraphs[0].add_run('– This flashes when there is an alarm state ')
    cell = table.cell(14, 1)
    cell.paragraphs[0].add_run(' ')
    cell = table.cell(15, 0)
    run = cell.paragraphs[0].add_run('Alarm – ')
    run.bold=True
    cell.paragraphs[0].add_run('This key will display any active alarms and a list of previous alarms since the last clear ')
    cell = table.cell(15, 1)
    cell.paragraphs[0].add_run(' ')
    cell = table.cell(16, 0)
    run = cell.paragraphs[0].add_run('Switching between modes matches the operator’s manual')
    cell = table.cell(16, 1)
    cell.paragraphs[0].add_run(' ')
    cell = table.cell(17, 0).merge(table.cell(17, 1))
    run = cell.paragraphs[0].add_run('Deviation / Deficiency Report No.')
    doc.add_paragraph()
    add_equipment_table(doc, equipment_number, document_number, creation_date)
    doc.add_paragraph()
    title = doc.add_heading('SECTION 9: TEST FUNCTIONS (Continued)', level=1)
    update_heading_style(title)
    table = doc.add_table(rows=9, cols=2)
    table.style = 'Table Grid' 
    table.allow_autofit = True
    cell = table.cell(0, 0).merge(table.cell(0, 1))
    run = cell.paragraphs[0].add_run('9.2 Test Function No. 2: Major Components Operating Parameters Verification ')
    run.bold = True
    cell = table.cell(1, 0).merge(table.cell(1, 1))
    run = cell.paragraphs[0].add_run('Objective: ')
    run.bold = True
    cell.paragraphs[0].add_run('To demonstrate that major components and / or their associated inputs / outputs to and from the controller operates as per design specifications. ')
    cell = table.cell(2, 0).merge(table.cell(2, 1))
    run = cell.paragraphs[0].add_run('Acceptance Criteria: ')
    run.bold = True
    cell.paragraphs[0].add_run('All features of the major components must operate as described in the Methodology Section')
    cell = table.cell(3, 0).merge(table.cell(3, 1))
    run = cell.paragraphs[0].add_run('Testing Equipment / Instruments: ')
    run.bold = True
    cell.paragraphs[0].add_run('Current Meter')
    cell = table.cell(4, 0).merge(table.cell(4, 1))
    run = cell.paragraphs[0].add_run('Methodology: ')
    run.bold = True
    cell.paragraphs[0].add_run('As follows.')
    cell = table.cell(5, 0)
    run = cell.paragraphs[0].add_run('Test')
    run.bold = True
    cell = table.cell(5, 1)
    run = cell.paragraphs[0].add_run('Conform (Y/N)')
    run.bold = True
    cell = table.cell(6, 0)
    run = cell.paragraphs[0].add_run('When the Trailer temperature is below the set point programmed in the temperature controller, the compressor solenoids are set for hot gas. ')
    cell = table.cell(6, 1)
    run = cell.paragraphs[0].add_run(' ')
    cell = table.cell(7, 0)
    run = cell.paragraphs[0].add_run('When the Trailer temperature is above the set point programmed in the temperature controller, the compressor solenoids are set for full cooling')
    cell = table.cell(7, 1)
    run = cell.paragraphs[0].add_run(' ')
    cell = table.cell(8, 0).merge(table.cell(8, 1))
    cell.paragraphs[0].add_run('Deviation / Deficiency Report No.')
    doc.add_paragraph()
    add_equipment_table(doc, equipment_number, document_number, creation_date)
    doc.add_paragraph()
    title = doc.add_heading('SECTION 9: TEST FUNCTIONS (Continued)', level=1)
    update_heading_style(title)
    table = doc.add_table(rows=9, cols=1)
    table.style = 'Table Grid' 
    table.allow_autofit = True
    cell = table.cell(0, 0)
    run = cell.paragraphs[0].add_run('9.3 Test Function No. 3: Trailer Temperature Distribution Verification ')
    cell = table.cell(1, 0)
    run = cell.paragraphs[0].add_run('Objective: ')
    run.bold = True
    cell.paragraphs[0].add_run('To demonstrate that the temperature is uniform and stable in accordance with the design specifications of the Trailer.')
    cell = table.cell(2, 0)
    run = cell.paragraphs[0].add_run('Acceptance Criteria: ')
    run.bold = True
    cell.paragraphs[0].add_run('For the six (6) hour empty Trailer mappings performed at (2°C to 8°C).and at (15°C to 25°C), the following criteria must be met: ')
    cell.paragraphs[0].add_run('\n\nThe temperature readings recorded by each of the sensors located in the Trailer must remain within (2°C to 8°C) and within (15°C to 25°C) respectively.')
    cell = table.cell(3, 0)
    run = cell.paragraphs[0].add_run('Testing Equipment/Instruments: ')
    run.bold = True
    cell.paragraphs[0].add_run('Temperature Loggers (8)')
    cell = table.cell(4, 0)
    run = cell.paragraphs[0].add_run('Methodology: ')
    run.bold = True
    cell.paragraphs[0].add_run('\n1.Position the loggers as per Appendix A, Drawing No. 1 – Temperature Mapping Locations. ')
    cell.paragraphs[0].add_run('\n2.Set the loggers to record the temperatures at 1-minute intervals.  ')
    cell.paragraphs[0].add_run('\n3.Verify the Trailer’s initial control set point.')
    cell.paragraphs[0].add_run('\n4.Once the Trailer’s temperature has stabilized, note the test start time and let it run for at least six (6) hours.')
    cell.paragraphs[0].add_run('\n5.Change and then verify the Trailer’s second control set-point.')
    cell.paragraphs[0].add_run('\n6.Once the Trailer’s temperature has stabilized, note the test start time and let it run for at least six (6) hours. ')
    cell.paragraphs[0].add_run('\n7.At the end of the test period, generate the raw data printouts and attach it in Appendix C, Item 3, Test Functions No. 3: Empty Trailer Temperature Distribution Verification –Data Logger Printouts. ')
    cell.paragraphs[0].add_run('\n8.The data of the test will be exported to spreadsheets used to identify the highest and lowest temperatures, as well as to calculate the average temperature for each time interval (1-minute intervals). The spreadsheets will be attached in Appendix B, Item 1, Test Function No. 3: Trailer Temperature Distribution Verification – Test Result Tables and Graphs.')
    doc.add_paragraph()
    add_equipment_table(doc, equipment_number, document_number, creation_date)
    doc.add_paragraph()
    title = doc.add_heading('SECTION 10: ANALYSIS AND CONCLUSIONS', level=1)
    update_heading_style(title)
    table = doc.add_table(rows=7, cols=1)
    table.style = 'Table Grid' 
    table.allow_autofit = True
    cell = table.cell(0, 0)
    run = cell.paragraphs[0].add_run('10.1 Test Function No. 1: Control Panel Verification')
    run.bold=True
    cell = table.cell(1, 0)
    run = cell.paragraphs[0].add_run('It has been demonstrated that the control panel features operate as per design specifications\n\nAll features of the control panel operate as described in the Methodology Section')
    cell = table.cell(2, 0)
    run = cell.paragraphs[0].add_run('10.2 Test Function No. 2: Major Components Operating Parameters Verification ')
    run.bold = True
    cell = table.cell(3, 0)
    cell.paragraphs[0].add_run('It has been demonstrated that the control panel features operate as per design specifications\n\nAll features of the control panel operate as described in the Methodology Section')
    cell = table.cell(4, 0)
    run = cell.paragraphs[0].add_run('10.3 Test Function No. 3: Trailer Temperature Distribution Verification')
    run.bold = True
    cell = table.cell(5, 0)
    cell.paragraphs[0].add_run('It has been demonstrated that the temperature is uniform and stable in accordance with the design specifications of the Trailer.\n\nFor the six (6) hours Trailer mappings performed at (2°C to 8°C). and (15°C to 25°C) the following criteria were met:\n\nThe temperature readings recorded by each of the sensors located in the Trailer must remain within (2°C to 8°C) and (15°C to 25°C) respectively. ')
    cell = table.cell(6, 0)
    run = cell.paragraphs[0].add_run('Deviation/Deficiency Report No.	            NAP ') 
    doc.add_paragraph()
    title=doc.add_heading('APPENDIX A: TRUCK/TRAILER IMAGE', level=1)
    update_heading_style(title)
    logo_width = Inches(4.5)
    image = doc.add_paragraph()
    run=image.add_run()
    image_url = "https://anvil.works/build/apps/PZEVDPC4ITSDLA7Q/code/assets/trailer.png"
    response = requests.get(image_url)

    if response.status_code == 200:
        
        try:
            # Open the image using PIL
            image = PilImage.open(BytesIO(response.content))
    
            # Convert and save it in a standard format
            image_stream = BytesIO()
            image.save(image_stream, format="PNG")  # Ensure PNG format
            image_stream.seek(0)  # Reset the stream position
    
            # Add image to document
            run.add_picture(image_stream, width=logo_width)  # Adjust width
    
            # Save document
            print("Document saved successfully!")
        except Exception as e:
            print(f"Error processing image: {e}")
    title=doc.add_heading('APPENDIX B: TEST RESULTS TABLES AND GRAPHS', level=1)
    update_heading_style(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.italic = True
    table = doc.add_table(rows=6, cols=1)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    cell = table.cell(0, 0)
    run = cell.paragraphs[0].add_run('Summary of Protocol')
    run.bold = True  # Make the text bold
    summary_protocol = ""
    if set_point == 20:
      summary_protocol = f"Skyline Cargo trailer {trailer_no} was mapped in {season} for temperature. The temperature range for the trailer being qualified is from 15°C to 25°C "
    else:
      summary_protocol = f"Skyline Cargo trailer {trailer_no} was mapped in {season} for temperature. The temperature range for the trailer being qualified is from 2°C to 8°C "
    table.cell(1, 0).text = summary_protocol
    cell = table.cell(2, 0)
    run = cell.paragraphs[0].add_run('Result from Study')
    run.bold = True  # Make the text bold
    cell = table.cell(3, 0)
    run = cell.paragraphs[0].add_run('IQ: ')
    run.bold = True
    cell.paragraphs[0].add_run('Verification of the Installation parameters was successfully recorded.\n\n')
    if start_input == 1:
        run = cell.paragraphs[0].add_run('OQ Empty Test Results:\n ')
    elif start_input == 2:
        run = cell.paragraphs[0].add_run('OQ Loaded Test Results:\n ')
    elif start_input == 3 or start_input == 4:
        run = cell.paragraphs[0].add_run('Loaded Trailer Power Failure and Open-Door Test:')
    elif start_input == 5:
        run = cell.paragraphs[0].add_run('Loaded Trailer Power Failure and Open-Door Test:')
    run.bold = True
    total_max, total_min, total_avg, total_median = process_all_files(files, start_datetime, end_datetime)
    if set_point == 20:
      cell.paragraphs[0].add_run(f'The temperature in the empty trailer was maintained for 6 hours within the temperature range from 15°C to 25°C, and ranged from a minimum {total_min}°C see data test sheets to a maximum of {total_max}°C see data test sheets.')
    if set_point == 5:
      cell.paragraphs[0].add_run(f'The temperature in the empty trailer was maintained for 6 hours within the temperature range from 2°C to 8°C, and ranged from a minimum {total_min}°C see data test sheets to a maximum of {total_max}°C see data test sheets.')
    cell = table.cell(4, 0)
    run = cell.paragraphs[0].add_run('Finding')
    run.bold = True  # Make the text bold
    if set_point == 20:
      finding = f"Skyline Cargo trailer {trailer_no} was successfully qualified for cold chain temperature specifications of 15°C to 25°C (Installation / Operational) in {season} condition."
    elif set_point == 5:
      finding = f"Skyline Cargo trailer {trailer_no} was successfully qualified for cold chain temperature specifications of 2°C to 8°C (Installation / Operational) in {season} condition."
    table.cell(5, 0).text = finding
    update_summary_table_style(table, 0)
    update_summary_table_style(table, 2)
    update_summary_table_style(table, 4)
    heading=doc.add_heading('Details', level=1)
    update_heading_style(heading)
    heading = doc.add_heading('1: Thresholds', level=2)
    update_heading_style(heading)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    update_style(table)
    table.cell(0, 0).text = 'Limit'
    table.cell(0, 1).text = 'Value'
    min_temp, max_temp = get_temp_range(files, set_point)
    table.cell(1, 0).text = 'Upper Temperature Limit'
    table.cell(1, 1).text = str(max_temp) + " °C"
    table.cell(2, 0).text = 'Lower Temperature Limit'
    table.cell(2, 1).text = str(min_temp) + " °C"
    heading = doc.add_heading('2: Conclusion', level=2)
    update_heading_style(heading)
    heading = doc.add_heading('2.1: Average/Minimum/Maximum Temperature', level=2)
    update_heading_style(heading)
    avg_temp, overall_min, overall_max, min_max_avg_df = process_file(doc, files, start_datetime, end_datetime)
    min_max_avg_plot_path=get_min_max_average_graph(min_max_avg_df)
    doc.add_heading('', level=1)
    heading=doc.add_heading('Temperature Mapping - Individual Sensor Results', level=3)
    update_heading_style(heading)
    doc.add_picture(min_max_avg_plot_path, width=Inches(6.0)) 
    doc.add_heading('', level=2)
    heading = doc.add_heading('2.2: Overall Temperature', level=2)
    update_heading_style(heading)
    table = doc.add_table(rows=2, cols=2)
    update_style(table)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    table.cell(0, 0).text = 'Measuring Point'
    table.cell(0, 1).text = 'Value[°C]'
    table.cell(1, 0).text = 'Overall average of all single average temperature values'
    avg_temp = round(avg_temp, 2)
    table.cell(1, 1).text = str(avg_temp)
    heading = doc.add_heading('2.3: Lowest/Highest Temperatures', level=2)
    update_heading_style(heading)
    low_high_calculator(files, doc, overall_min, 'Lowest', start_datetime, end_datetime)
    doc.add_heading('', level=2)
    low_high_calculator(files, doc, overall_max, 'Highest', start_datetime, end_datetime)
    heading = doc.add_heading('3: Mapping Results', level=2)
    update_heading_style(heading)
    combine_image_path, number_of_files = get_combined_graph(files, start_datetime, end_datetime)
    heading=doc.add_heading('3.1: Overview', level=2)
    update_heading_style(heading)
    doc.add_picture(combine_image_path, width=Inches(6.0)) 
    doc.add_heading('', level=2)
    table = doc.add_table(rows=6, cols=2)
    update_style(table)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    table.cell(0, 0).text = 'Statistics'
    table.cell(0, 1).text = 'Value'
    table.cell(1, 0).text = 'Range'
    table.cell(1, 1).text = str(start_datetime) + '-' + str(end_datetime)
    total_max, total_min, total_avg, total_median = process_all_files(files, start_datetime, end_datetime)
    table.cell(2, 0).text = 'Average'
    table.cell(2, 1).text = str(total_avg) + ' °C'
    table.cell(3, 0).text = 'Min/Max Thresholds'
    table.cell(3, 1).text = str(total_max) + ' / ' + str(total_min) +  ' °C'
    table.cell(4, 0).text = 'Number of files'
    table.cell(4, 1).text = str(number_of_files)
    table.cell(5, 0).text = 'Number of sensors'
    table.cell(5, 1).text = str(number_of_files)
    read_and_filter_data(doc, files,start_datetime, end_datetime)
    doc.add_heading(' ', level=2)
    title=doc.add_heading('4. Analysis and Conclusion', level=2)
    update_heading_style(title)
    table = doc.add_table(rows=2, cols=1)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    cell = table.cell(0, 0)
    run = cell.paragraphs[0].add_run('Test Function No. 1: Temperature Distribution Verification')
    run.bold = True  # Make the text bold
    summary_protocol = "It has been demonstrated that the temperature is uniform and stable in accordance with the design specifications of the Trailer.\n\n"
    summary_protocol += "For the "
    if start_input == 1:
        summary_protocol += '6-Hour Mapping-Empty Trailer'
    elif start_input == 2:
        summary_protocol += '6-Hour Mapping-Loaded Trailer'
    elif start_input == 3:
        summary_protocol += '1-Hour Temperature-Control Power Failure-Loaded Trailer'
    elif start_input == 4:
        summary_protocol += '1-Minute Open Door-Loaded Trailer'
    else:
        summary_protocol += '2-Hour Field Shipment Test-Loaded Trailer'
    summary_protocol += ' performed at '
    if set_point == 5:
      summary_protocol += "(2°C to 8°C)"
    elif set_point == 20:
      summary_protocol += "(15°C to 25°C)"
    summary_protocol += '.The following criteria must be met: '
    summary_protocol += "\n\nThe temperature readings recorded by each of the sensors located in the trailer must remain within"
    if set_point == 5:
      summary_protocol += "(2°C to 8°C)"
    elif set_point == 20:
      summary_protocol += "(15°C to 25°C)"
    summary_protocol += '\n\n                                                  Min             Median               Max'
    summary_protocol += '\nTemperature                        '
    summary_protocol += f'{total_min} °C       '
    summary_protocol += f'{total_median} °C              '
    summary_protocol += f'{total_max} °C   ' 
    summary_protocol += '\nDifference from Setpoint       '
    diff_temp=float(set_point)-float(total_min)
    diff_temp=round(diff_temp, 1)
    summary_protocol += f'{diff_temp} °C       '
    diff_temp=float(set_point)-float(total_median)
    diff_temp=round(diff_temp, 1)
    summary_protocol += f'{diff_temp} °C              '
    diff_temp=float(set_point)-float(total_max)
    diff_temp=round(diff_temp, 1)
    summary_protocol += f'{diff_temp} °C   ' 
    table.cell(1, 0).text = summary_protocol
    update_summary_table_style(table, 0)
    doc_stream = io.BytesIO()
    doc.save(doc_stream) 
    doc_stream.seek(0)
    return BlobMedia(
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
        doc_stream.read(), 
        'Temperature Mapping Protocol.docx'
    ) 
@anvil.server.callable
def save_user_choice(start_digit, author_name, start_date, start_time, end_date, end_time, temp_value, application_name, company_name, files, trailer_no, season, protocol_number, document_number, make,model_number,vin_number, equipment_number, operating_conditions, creation_date, equipment_type, revision_date):
    start_date_str = start_date.strftime('%Y-%m-%d')
    start_str = start_date_str + " " + start_time
    start_datetime = pd.to_datetime(start_str)
    end_date_str = end_date.strftime('%Y-%m-%d')
    end_str = end_date_str + " " + end_time
    end_datetime = pd.to_datetime(end_str)
    print("start_datetime=", start_datetime)
    print("end_datetime=", end_datetime)
    return create_document(files, start_datetime, end_datetime, start_digit, temp_value, company_name, author_name, application_name, trailer_no, season, protocol_number, document_number, make,model_number,vin_number, equipment_number, operating_conditions, creation_date, equipment_type, revision_date)
          
